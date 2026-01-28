"""Document ingestion service."""
import os
import uuid
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from docx import Document
import pdfplumber
from sqlalchemy.orm import Session

from app.models.toolkit import ToolkitDocument, ToolkitChunk

logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse PDF file and extract text content.

    Args:
        file_path: Path to PDF file

    Returns:
        List of content blocks with text and metadata
    """
    content_blocks = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                # Split by paragraphs (double newlines)
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        content_blocks.append({
                            'type': 'paragraph',
                            'text': para,
                            'heading': f'Page {page_num}',
                            'page': page_num
                        })

    return content_blocks


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse DOCX file and extract structured content.

    Args:
        file_path: Path to DOCX file

    Returns:
        List of content blocks with text and metadata
    """
    doc = Document(file_path)
    content_blocks = []
    current_heading = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph is a heading
        if para.style.name.startswith('Heading'):
            current_heading = text
            content_blocks.append({
                'type': 'heading',
                'text': text,
                'level': para.style.name,
                'heading': current_heading
            })
        else:
            content_blocks.append({
                'type': 'paragraph',
                'text': text,
                'heading': current_heading
            })

    return content_blocks


def chunk_content(
    content_blocks: List[Dict[str, Any]],
    target_size: int = 1000,
    overlap: int = 150
) -> List[Dict[str, Any]]:
    """
    Chunk content into manageable pieces with overlap.

    Args:
        content_blocks: List of content blocks from parse_docx
        target_size: Target chunk size in characters (800-1200)
        overlap: Overlap between chunks in characters

    Returns:
        List of chunks with text and metadata
    """
    chunks = []
    current_chunk = []
    current_size = 0
    current_heading = None
    chunk_index = 0

    for block in content_blocks:
        if block['type'] == 'heading':
            current_heading = block['text']

        text = block['text']
        text_len = len(text)

        # If adding this block would exceed target size and we have content, create chunk
        if current_size + text_len > target_size and current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'chunk_text': chunk_text,
                'chunk_index': chunk_index,
                'heading': current_heading,
                'metadata': {'char_count': len(chunk_text)}
            })
            chunk_index += 1

            # Keep last part for overlap
            if current_chunk:
                overlap_text = current_chunk[-1]
                if len(overlap_text) > overlap:
                    overlap_text = overlap_text[-overlap:]
                current_chunk = [overlap_text, text]
                current_size = len(overlap_text) + text_len
            else:
                current_chunk = [text]
                current_size = text_len
        else:
            current_chunk.append(text)
            current_size += text_len

    # Add remaining content as final chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        chunks.append({
            'chunk_text': chunk_text,
            'chunk_index': chunk_index,
            'heading': current_heading,
            'metadata': {'char_count': len(chunk_text)}
        })

    return chunks


def ingest_document(
    db: Session,
    file_path: str,
    version_tag: str,
    source_filename: str,
    create_embeddings: bool = True
) -> ToolkitDocument:
    """
    Ingest a document: parse, chunk, and store in database.

    Args:
        db: Database session
        file_path: Path to uploaded DOCX or PDF file
        version_tag: Version identifier
        source_filename: Original filename
        create_embeddings: Whether to create embeddings (requires OpenAI API)

    Returns:
        Created ToolkitDocument instance
    """
    # Check if version already exists
    existing = db.query(ToolkitDocument).filter(
        ToolkitDocument.version_tag == version_tag
    ).first()
    if existing:
        raise ValueError(f"Version tag '{version_tag}' already exists")

    # Parse document based on file type
    if file_path.lower().endswith('.pdf'):
        content_blocks = parse_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        content_blocks = parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type. Must be .docx or .pdf")

    # Create chunks
    chunks = chunk_content(content_blocks)

    # Create document record
    doc = ToolkitDocument(
        version_tag=version_tag,
        source_filename=source_filename,
        file_path=file_path,
        chunk_count=len(chunks)
    )
    db.add(doc)
    db.flush()  # Get document ID

    # Create chunk records
    chunk_objects = []
    for chunk_data in chunks:
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data['chunk_text'],
            chunk_index=chunk_data['chunk_index'],
            heading=chunk_data.get('heading'),
            chunk_metadata=chunk_data.get('metadata'),
            embedding=None  # Will be populated by embeddings service
        )
        chunk_objects.append(chunk)

    db.bulk_save_objects(chunk_objects)
    db.commit()
    db.refresh(doc)

    # Create embeddings if requested
    if create_embeddings:
        from app.services.embeddings import create_embeddings_for_document
        create_embeddings_for_document(db, doc.id)

    return doc


def reindex_document(db: Session, document_id: str) -> ToolkitDocument:
    """
    Reindex a document: re-run chunking and embeddings.

    This deletes existing chunks and recreates them from the source file.

    Args:
        db: Database session
        document_id: ID of document to reindex

    Returns:
        Updated ToolkitDocument instance

    Raises:
        ValueError: If document not found or file doesn't exist
    """
    # Get document
    doc = db.query(ToolkitDocument).filter(ToolkitDocument.id == document_id).first()

    if not doc:
        raise ValueError(f"Document {document_id} not found")

    if not os.path.exists(doc.file_path):
        raise ValueError(f"Source file not found: {doc.file_path}")

    # Delete existing chunks
    db.query(ToolkitChunk).filter(ToolkitChunk.document_id == document_id).delete()
    db.commit()

    # Re-parse document based on file type
    if doc.file_path.lower().endswith('.pdf'):
        content_blocks = parse_pdf(doc.file_path)
    elif doc.file_path.lower().endswith('.docx'):
        content_blocks = parse_docx(doc.file_path)
    else:
        raise ValueError(f"Unsupported file type: {doc.file_path}")

    # Create new chunks
    chunks = chunk_content(content_blocks)

    # Update chunk count
    doc.chunk_count = len(chunks)

    # Create chunk records
    chunk_objects = []
    for chunk_data in chunks:
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data['chunk_text'],
            chunk_index=chunk_data['chunk_index'],
            heading=chunk_data.get('heading'),
            chunk_metadata=chunk_data.get('metadata'),
            embedding=None  # Will be populated by embeddings service
        )
        chunk_objects.append(chunk)

    db.bulk_save_objects(chunk_objects)
    db.commit()
    db.refresh(doc)

    # Recreate embeddings
    from app.services.embeddings import create_embeddings_for_document
    create_embeddings_for_document(db, doc.id)

    return doc


def ingest_from_kit(
    db: Session,
    version_tag: str = "kit-v1",
    create_embeddings: bool = True
) -> ToolkitDocument:
    """
    Ingest toolkit content from /kit JSON files into the database.

    Reads structured tool, cluster, and foundation JSON files from the kit
    directory and creates chunked records with enriched metadata for RAG search.

    Args:
        db: Database session
        version_tag: Version identifier for this ingestion
        create_embeddings: Whether to create embeddings

    Returns:
        Created ToolkitDocument instance
    """
    from app.services.kit_loader import (
        get_all_tools, get_all_clusters, get_all_foundations, get_all_sources, clear_cache
    )

    # Clear cache to pick up any new data
    clear_cache()

    # Check if version already exists
    existing = db.query(ToolkitDocument).filter(
        ToolkitDocument.version_tag == version_tag
    ).first()
    if existing:
        # Deactivate old version
        existing.is_active = False
        db.commit()
        # Bump version tag
        version_tag = f"{version_tag}-{uuid.uuid4().hex[:8]}"

    tools = get_all_tools()
    clusters = get_all_clusters()
    foundations = get_all_foundations()
    sources_data = get_all_sources()
    source_entries = sources_data.get("entries", [])

    logger.info(f"Ingesting from kit: {len(tools)} tools, {len(clusters)} clusters, {len(foundations)} foundations, {len(source_entries)} sources")

    # Build chunks from all kit content
    all_chunks = []
    chunk_index = 0

    # 1. Chunk each tool's content
    for tool in tools:
        # Build the full text for this tool
        parts = [f"Tool: {tool['name']}"]
        if tool.get("description"):
            parts.append(tool["description"])
        if tool.get("purpose"):
            parts.append(f"Purpose: {tool['purpose']}")
        if tool.get("journalism_relevance"):
            parts.append(f"Journalism Relevance: {tool['journalism_relevance']}")
        if tool.get("comments"):
            parts.append(f"Comments: {tool['comments']}")

        cdi = tool.get("cdi_scores", {})
        parts.append(f"CDI Score - Cost: {cdi.get('cost', 0)}/10, Difficulty: {cdi.get('difficulty', 0)}/10, Invasiveness: {cdi.get('invasiveness', 0)}/10")

        td = tool.get("time_dividend", {})
        if td.get("time_saved"):
            parts.append(f"Time Saved: {td['time_saved']}")
        if td.get("reinvestment"):
            parts.append(f"Time Reinvestment: {td['reinvestment']}")

        full_text = "\n\n".join(parts)

        # Metadata for this tool's chunks
        metadata = {
            "type": "tool",
            "tool_name": tool["name"],
            "tool_slug": tool["slug"],
            "cluster": tool.get("cluster_name", ""),
            "cluster_slug": tool.get("cluster_slug", ""),
            "cdi_scores": cdi,
            "tags": tool.get("tags", []),
        }

        # Create chunks from tool text
        tool_chunks = chunk_content(
            [{"type": "paragraph", "text": full_text, "heading": tool["name"]}],
            target_size=1000,
            overlap=150
        )

        for tc in tool_chunks:
            tc["metadata"] = {**tc.get("metadata", {}), **metadata}
            tc["heading"] = tool["name"]
            tc["chunk_index"] = chunk_index
            chunk_index += 1
            all_chunks.append(tc)

    # 2. Chunk each cluster's description and teaching content
    for cluster in clusters:
        parts = [f"Cluster: {cluster['name']}"]
        if cluster.get("description"):
            parts.append(cluster["description"])
        if cluster.get("teaching_guidance"):
            parts.append(f"Teaching Guidance: {cluster['teaching_guidance']}")
        if cluster.get("where_to_start"):
            parts.append(f"Where to Start: {cluster['where_to_start']}")
        for ex in cluster.get("exercises", []):
            parts.append(f"Exercise {ex.get('number', '')}: {ex.get('title', '')} - {ex.get('description', '')}")

        full_text = "\n\n".join(parts)
        metadata = {
            "type": "cluster",
            "cluster": cluster["name"],
            "cluster_slug": cluster["slug"],
        }

        cluster_chunks = chunk_content(
            [{"type": "paragraph", "text": full_text, "heading": cluster["name"]}],
            target_size=1000,
            overlap=150
        )

        for cc in cluster_chunks:
            cc["metadata"] = {**cc.get("metadata", {}), **metadata}
            cc["heading"] = cluster["name"]
            cc["chunk_index"] = chunk_index
            chunk_index += 1
            all_chunks.append(cc)

    # 3. Chunk foundational content
    for foundation in foundations:
        content = foundation.get("content", "")
        if not content:
            continue

        title = foundation.get("title", "Foundation")
        metadata = {
            "type": "foundation",
            "foundation_slug": foundation.get("slug", ""),
        }

        foundation_chunks = chunk_content(
            [{"type": "paragraph", "text": content, "heading": title}],
            target_size=1000,
            overlap=150
        )

        for fc in foundation_chunks:
            fc["metadata"] = {**fc.get("metadata", {}), **metadata}
            fc["heading"] = title
            fc["chunk_index"] = chunk_index
            chunk_index += 1
            all_chunks.append(fc)

    # 4. Chunk source citations
    for entry in source_entries:
        parts = [f"Source: {entry.get('title', '')}"]
        if entry.get("excerpt"):
            parts.append(f"Key excerpt: {entry['excerpt']}")
        if entry.get("why_it_matters"):
            parts.append(f"Why it matters: {entry['why_it_matters']}")
        if entry.get("ai_extract"):
            parts.append(entry["ai_extract"])
        if entry.get("url"):
            parts.append(f"URL: {entry['url']}")

        full_text = "\n\n".join(parts)
        metadata = {
            "type": "source",
            "source_id": entry.get("entry_id", ""),
            "batch": entry.get("batch", 0),
            "theme": entry.get("theme", ""),
            "url": entry.get("url", ""),
        }

        source_chunks = chunk_content(
            [{"type": "paragraph", "text": full_text, "heading": entry.get("title", "Source")}],
            target_size=1000,
            overlap=150
        )

        for sc in source_chunks:
            sc["metadata"] = {**sc.get("metadata", {}), **metadata}
            sc["heading"] = entry.get("title", "Source")
            sc["chunk_index"] = chunk_index
            chunk_index += 1
            all_chunks.append(sc)

    # Create document record
    from pathlib import Path
    kit_dir = Path(__file__).resolve().parent.parent.parent / "kit"

    doc = ToolkitDocument(
        version_tag=version_tag,
        source_filename="toolkit.pdf",
        file_path=str(kit_dir / "toolkit.pdf"),
        chunk_count=len(all_chunks)
    )
    db.add(doc)
    db.flush()

    # Create chunk records
    chunk_objects = []
    for chunk_data in all_chunks:
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data["chunk_text"],
            chunk_index=chunk_data["chunk_index"],
            heading=chunk_data.get("heading"),
            chunk_metadata=chunk_data.get("metadata"),
            embedding=None
        )
        chunk_objects.append(chunk)

    db.bulk_save_objects(chunk_objects)
    db.commit()
    db.refresh(doc)

    logger.info(f"Ingested {len(all_chunks)} chunks from kit (version: {version_tag})")

    # Create embeddings if requested
    if create_embeddings:
        from app.services.embeddings import create_embeddings_for_document
        create_embeddings_for_document(db, doc.id)

    return doc


def ingest_batch_pdfs(
    db: Session,
    create_embeddings: bool = True
) -> List[ToolkitDocument]:
    """
    Ingest all batch PDF files from the /kit directory.

    Finds all batch*.pdf files and ingests each as a separate document.

    Args:
        db: Database session
        create_embeddings: Whether to create embeddings

    Returns:
        List of created ToolkitDocument instances
    """
    kit_dir = Path(__file__).resolve().parent.parent.parent / "kit"

    # Find all batch PDFs
    batch_pdfs = sorted(kit_dir.glob("batch*.pdf"), key=lambda p: int(p.stem.replace("batch", "") or 0))

    if not batch_pdfs:
        logger.warning("No batch PDFs found in kit directory")
        return []

    logger.info(f"Found {len(batch_pdfs)} batch PDFs to ingest")

    ingested_docs = []

    for pdf_path in batch_pdfs:
        batch_name = pdf_path.stem  # e.g., "batch1"
        version_tag = f"batch-{batch_name}"

        # Check if already exists
        existing = db.query(ToolkitDocument).filter(
            ToolkitDocument.version_tag == version_tag
        ).first()

        if existing:
            logger.info(f"Skipping {batch_name} - already ingested (version: {version_tag})")
            ingested_docs.append(existing)
            continue

        try:
            logger.info(f"Ingesting {pdf_path.name}...")

            # Parse PDF
            content_blocks = parse_pdf(str(pdf_path))

            if not content_blocks:
                logger.warning(f"No content extracted from {pdf_path.name}")
                continue

            # Create chunks
            chunks = chunk_content(content_blocks)

            # Create document record
            doc = ToolkitDocument(
                version_tag=version_tag,
                source_filename=pdf_path.name,
                file_path=str(pdf_path),
                chunk_count=len(chunks)
            )
            db.add(doc)
            db.flush()

            # Create chunk records
            chunk_objects = []
            for chunk_data in chunks:
                # Add batch metadata
                chunk_metadata = chunk_data.get('metadata', {})
                chunk_metadata['batch'] = batch_name
                chunk_metadata['type'] = 'source_pdf'

                chunk = ToolkitChunk(
                    document_id=doc.id,
                    chunk_text=chunk_data['chunk_text'],
                    chunk_index=chunk_data['chunk_index'],
                    heading=chunk_data.get('heading'),
                    chunk_metadata=chunk_metadata,
                    embedding=None
                )
                chunk_objects.append(chunk)

            db.bulk_save_objects(chunk_objects)
            db.commit()
            db.refresh(doc)

            logger.info(f"Ingested {pdf_path.name}: {len(chunks)} chunks")

            # Create embeddings if requested
            if create_embeddings:
                from app.services.embeddings import create_embeddings_for_document
                create_embeddings_for_document(db, doc.id)

            ingested_docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to ingest {pdf_path.name}: {e}")
            db.rollback()
            continue

    logger.info(f"Batch PDF ingestion complete: {len(ingested_docs)} documents")
    return ingested_docs
