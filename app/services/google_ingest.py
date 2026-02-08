"""Google Drive file listing, export, and ingestion pipeline."""
import hashlib
import logging
import re
import uuid
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional

from bs4 import BeautifulSoup
from sqlalchemy.orm import Session

from app.models.google_drive import GoogleConnection, GoogleSyncItem
from app.models.library_item import LibraryItem
from app.models.toolkit import ToolkitDocument, ToolkitChunk
from app.services.google_auth import get_drive_service
from app.services.ingestion import chunk_content

logger = logging.getLogger(__name__)


# ── Drive file listing ──────────────────────────────────────────────────────

def list_drive_files(
    db: Session,
    connection: GoogleConnection,
    folder_id: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """List files and folders in a Google Drive folder.

    Args:
        db: Database session (needed for token refresh)
        connection: Google OAuth connection
        folder_id: Drive folder ID, or None for root

    Returns:
        List of dicts with id, name, mimeType, modifiedTime, parents
    """
    service = get_drive_service(db, connection)

    query_parts = ["trashed = false"]
    if folder_id:
        query_parts.append(f"'{folder_id}' in parents")
    else:
        query_parts.append("'root' in parents")

    q = " and ".join(query_parts)

    results = []
    page_token = None

    while True:
        resp = service.files().list(
            q=q,
            pageSize=100,
            fields="nextPageToken, files(id, name, mimeType, modifiedTime, parents)",
            orderBy="folder,name",
            pageToken=page_token,
        ).execute()

        results.extend(resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    return results


# ── Content export ──────────────────────────────────────────────────────────

def _html_to_markdown(html: str) -> str:
    """Convert HTML to simple markdown using BeautifulSoup."""
    soup = BeautifulSoup(html, "html.parser")

    # Remove style and script tags
    for tag in soup.find_all(["style", "script"]):
        tag.decompose()

    lines = []

    for element in soup.descendants:
        if element.name and element.string:
            text = element.string.strip()
            if not text:
                continue

            if element.name in ("h1",):
                lines.append(f"\n# {text}\n")
            elif element.name in ("h2",):
                lines.append(f"\n## {text}\n")
            elif element.name in ("h3",):
                lines.append(f"\n### {text}\n")
            elif element.name in ("h4",):
                lines.append(f"\n#### {text}\n")
            elif element.name in ("h5",):
                lines.append(f"\n##### {text}\n")
            elif element.name in ("h6",):
                lines.append(f"\n###### {text}\n")
            elif element.name == "li":
                lines.append(f"- {text}")
            elif element.name == "a":
                href = element.get("href", "")
                if href:
                    lines.append(f"[{text}]({href})")
            elif element.name in ("p", "div", "span", "td", "th"):
                lines.append(text)
        elif element.name == "br":
            lines.append("")

    # Clean up: collapse multiple blank lines
    md = "\n".join(lines)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def export_google_doc(db: Session, connection: GoogleConnection, file_id: str) -> str:
    """Export a Google Doc as HTML and convert to markdown."""
    service = get_drive_service(db, connection)
    html = service.files().export(
        fileId=file_id, mimeType="text/html"
    ).execute()

    if isinstance(html, bytes):
        html = html.decode("utf-8")

    return _html_to_markdown(html)


def export_google_sheet(db: Session, connection: GoogleConnection, file_id: str) -> str:
    """Export a Google Sheet as CSV text."""
    service = get_drive_service(db, connection)
    csv_bytes = service.files().export(
        fileId=file_id, mimeType="text/csv"
    ).execute()

    if isinstance(csv_bytes, bytes):
        return csv_bytes.decode("utf-8")
    return csv_bytes


def export_file_content(
    db: Session, connection: GoogleConnection, file_id: str, mime_type: str
) -> Optional[str]:
    """Export file content based on mime type. Returns markdown text or None."""
    if mime_type == "application/vnd.google-apps.document":
        return export_google_doc(db, connection, file_id)

    if mime_type == "application/vnd.google-apps.spreadsheet":
        return export_google_sheet(db, connection, file_id)

    # Binary files: download and extract text
    if mime_type == "application/pdf":
        return _extract_pdf(db, connection, file_id)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(db, connection, file_id)

    logger.warning(f"Unsupported mime type for export: {mime_type}")
    return None


def _extract_pdf(db: Session, connection: GoogleConnection, file_id: str) -> str:
    """Download a PDF from Drive and extract text via pdfplumber."""
    import tempfile
    import pdfplumber

    service = get_drive_service(db, connection)
    content = service.files().get_media(fileId=file_id).execute()

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()

        pages = []
        with pdfplumber.open(tmp.name) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

    return "\n\n".join(pages)


def _extract_docx(db: Session, connection: GoogleConnection, file_id: str) -> str:
    """Download a DOCX from Drive and extract text via python-docx."""
    import tempfile
    from docx import Document

    service = get_drive_service(db, connection)
    content = service.files().get_media(fileId=file_id).execute()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()

        doc = Document(tmp.name)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    return "\n\n".join(paragraphs)


# ── Ingestion pipeline ──────────────────────────────────────────────────────

def _compute_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def ingest_to_library(
    db: Session,
    sync_item: GoogleSyncItem,
    markdown_content: str,
) -> None:
    """Ingest content into the public library as a LibraryItem + ToolkitChunks."""
    title = sync_item.google_file_name
    # Strip common extensions from the title
    for ext in (".gdoc", ".docx", ".pdf", ".gsheet"):
        if title.lower().endswith(ext):
            title = title[: -len(ext)]

    source_id = f"gdrive:{sync_item.google_file_id}"

    # Check for existing library item by source_id
    existing = db.query(LibraryItem).filter(LibraryItem.source_id == source_id).first()

    if existing:
        existing.content_markdown = markdown_content
        existing.content_hash = _compute_hash(markdown_content)
        existing.title = title
        existing.slug = LibraryItem.generate_slug(title)
        db.flush()
        library_item = existing
    else:
        slug = LibraryItem.generate_slug(title)
        # Ensure slug uniqueness
        base_slug = slug
        counter = 1
        while db.query(LibraryItem).filter(LibraryItem.slug == slug).first():
            slug = f"{base_slug}-{counter}"
            counter += 1

        library_item = LibraryItem(
            title=title,
            slug=slug,
            document_type="guidance",
            content_markdown=markdown_content,
            is_published=False,
            source_id=source_id,
            content_hash=_compute_hash(markdown_content),
        )
        db.add(library_item)
        db.flush()

    sync_item.library_item_id = library_item.id

    # Create toolkit chunks for RAG search
    _create_chunks_for_content(
        db,
        markdown_content,
        title,
        version_tag=f"gdrive:library:{sync_item.google_file_id}",
        metadata_extra={"type": "library", "source": "google_drive"},
        sync_item=sync_item,
    )


def ingest_to_organization(
    db: Session,
    sync_item: GoogleSyncItem,
    markdown_content: str,
) -> None:
    """Ingest content into an organization's knowledge base as ToolkitChunks."""
    org_id = str(sync_item.target_id)
    title = sync_item.google_file_name

    _create_chunks_for_content(
        db,
        markdown_content,
        title,
        version_tag=f"org:{org_id}:{sync_item.google_file_id}",
        metadata_extra={
            "type": "organization",
            "organization_id": org_id,
            "source": "google_drive",
        },
        sync_item=sync_item,
    )


def _create_chunks_for_content(
    db: Session,
    markdown_content: str,
    title: str,
    version_tag: str,
    metadata_extra: dict,
    sync_item: GoogleSyncItem,
) -> None:
    """Create ToolkitDocument + ToolkitChunks from markdown content."""
    # Check for existing document with this version_tag
    existing_doc = db.query(ToolkitDocument).filter(
        ToolkitDocument.version_tag == version_tag
    ).first()

    if existing_doc:
        # Delete old chunks, reuse document
        db.query(ToolkitChunk).filter(ToolkitChunk.document_id == existing_doc.id).delete()
        doc = existing_doc
        doc.is_active = True
        doc.is_ingested = True
    else:
        doc = ToolkitDocument(
            version_tag=version_tag,
            source_filename=sync_item.google_file_name,
            file_path=f"gdrive://{sync_item.google_file_id}",
            chunk_count=0,
            is_ingested=True,
        )
        db.add(doc)
        db.flush()

    # Build content blocks and chunk
    content_blocks = [
        {"type": "paragraph", "text": markdown_content, "heading": title}
    ]
    chunks = chunk_content(content_blocks, target_size=1000, overlap=150)

    chunk_objects = []
    for chunk_data in chunks:
        chunk_meta = {**chunk_data.get("metadata", {}), **metadata_extra}
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data["chunk_text"],
            chunk_index=chunk_data["chunk_index"],
            heading=chunk_data.get("heading"),
            chunk_metadata=chunk_meta,
            embedding=None,
        )
        chunk_objects.append(chunk)

    doc.chunk_count = len(chunk_objects)
    db.bulk_save_objects(chunk_objects)
    db.flush()

    sync_item.toolkit_document_id = doc.id

    # Create embeddings
    try:
        from app.services.embeddings import create_embeddings_for_document
        create_embeddings_for_document(db, doc.id)
    except Exception as e:
        logger.warning(f"Failed to create embeddings for {version_tag}: {e}")


def ingest_google_file(
    db: Session,
    sync_item: GoogleSyncItem,
    connection: GoogleConnection,
) -> None:
    """Full ingestion pipeline for a single Google Drive file."""
    sync_item.sync_status = "syncing"
    sync_item.error_message = None
    db.commit()

    try:
        # Export content
        content = export_file_content(
            db, connection, sync_item.google_file_id, sync_item.google_mime_type
        )

        if content is None:
            sync_item.sync_status = "error"
            sync_item.error_message = f"Unsupported file type: {sync_item.google_mime_type}"
            db.commit()
            return

        content_hash = _compute_hash(content)

        # Skip if content unchanged
        if sync_item.content_hash == content_hash and sync_item.sync_status == "synced":
            logger.info(f"Content unchanged for {sync_item.google_file_name}, skipping")
            return

        # Route to the right target
        if sync_item.target_type == "library":
            ingest_to_library(db, sync_item, content)
        elif sync_item.target_type == "organization":
            ingest_to_organization(db, sync_item, content)

        sync_item.sync_status = "synced"
        sync_item.content_hash = content_hash
        sync_item.last_synced_at = datetime.now(timezone.utc)
        db.commit()

        logger.info(f"Ingested Google file: {sync_item.google_file_name} → {sync_item.target_type}")

    except Exception as e:
        db.rollback()
        sync_item.sync_status = "error"
        sync_item.error_message = str(e)[:500]
        db.commit()
        logger.error(f"Failed to ingest {sync_item.google_file_name}: {e}")


def sync_all_items(db: Session, connection: GoogleConnection) -> dict:
    """Re-sync all items for a connection. Returns summary stats."""
    items = db.query(GoogleSyncItem).filter(
        GoogleSyncItem.connection_id == connection.id
    ).all()

    stats = {"total": len(items), "synced": 0, "errors": 0, "skipped": 0}

    for item in items:
        old_hash = item.content_hash
        ingest_google_file(db, item, connection)

        if item.sync_status == "synced":
            if item.content_hash == old_hash:
                stats["skipped"] += 1
            else:
                stats["synced"] += 1
        elif item.sync_status == "error":
            stats["errors"] += 1

    return stats
