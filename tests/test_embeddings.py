"""Embedding provider tests."""
import pytest
import os
import tempfile
from docx import Document

from app.models.toolkit import ToolkitChunk
from app.services.embeddings import (
    OpenAIEmbeddingProvider,
    LocalStubEmbeddingProvider,
    get_embedding_provider,
    create_embedding,
    create_embeddings_for_document
)
from app.services.ingestion import ingest_document
from app.settings import settings


def test_local_stub_provider_deterministic():
    """Test local stub provider creates deterministic embeddings."""
    provider = LocalStubEmbeddingProvider(dimensions=1536)

    text = "Test content for embedding"

    # Same text should produce same embedding
    embedding1 = provider.create_embedding(text)
    embedding2 = provider.create_embedding(text)

    assert embedding1 == embedding2
    assert len(embedding1) == 1536
    assert len(embedding2) == 1536

    # Different text should produce different embeddings
    embedding3 = provider.create_embedding("Different text")
    assert embedding3 != embedding1


def test_local_stub_provider_normalized():
    """Test local stub provider creates normalized embeddings."""
    provider = LocalStubEmbeddingProvider(dimensions=1536)

    embedding = provider.create_embedding("Test content")

    # Calculate magnitude
    magnitude = sum(x**2 for x in embedding) ** 0.5

    # Should be normalized (magnitude close to 1.0)
    assert abs(magnitude - 1.0) < 0.0001


def test_local_stub_provider_dimensions():
    """Test local stub provider respects custom dimensions."""
    provider = LocalStubEmbeddingProvider(dimensions=512)

    embedding = provider.create_embedding("Test content")

    assert len(embedding) == 512
    assert provider.dimensions == 512


def test_get_embedding_provider_local_stub(monkeypatch):
    """Test getting local stub provider from settings."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    provider = get_embedding_provider()

    assert provider is not None
    assert isinstance(provider, LocalStubEmbeddingProvider)
    assert provider.dimensions == settings.EMBEDDING_DIMENSIONS


def test_get_embedding_provider_openai_requires_key(monkeypatch):
    """Test OpenAI provider fails without API key."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "openai")
    monkeypatch.setattr(settings, "OPENAI_API_KEY", "sk-your-api-key-here")

    with pytest.raises(ValueError, match="OPENAI_API_KEY is not configured"):
        get_embedding_provider()


def test_create_embedding_with_local_stub():
    """Test create_embedding function with local stub provider."""
    provider = LocalStubEmbeddingProvider(dimensions=1536)

    embedding = create_embedding("Test content", provider=provider)

    assert embedding is not None
    assert len(embedding) == 1536


def test_create_embeddings_for_document(db_session, monkeypatch):
    """Test creating embeddings for all chunks in a document."""
    # Use local_stub provider
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create test document
    doc = Document()
    doc.add_heading('Test Heading', level=1)
    doc.add_paragraph('This is a test paragraph with some content.')
    doc.add_paragraph('This is another paragraph with more content.')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    try:
        # Ingest document without embeddings first
        document = ingest_document(
            db=db_session,
            file_path=temp_file.name,
            version_tag="embedding-test",
            source_filename="test.docx",
            create_embeddings=False
        )

        # Verify no embeddings yet
        chunks_before = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).all()

        assert all(chunk.embedding is None for chunk in chunks_before)

        # Create embeddings
        count = create_embeddings_for_document(db_session, document.id)

        # Should have created embeddings for all chunks
        assert count == document.chunk_count
        assert count > 0

        # Verify embeddings were added
        chunks_after = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).all()

        for chunk in chunks_after:
            assert chunk.embedding is not None
            assert len(chunk.embedding) == settings.EMBEDDING_DIMENSIONS

    finally:
        os.unlink(temp_file.name)


def test_ingest_with_embeddings_local_stub(db_session, monkeypatch):
    """Test document ingestion with embeddings using local stub."""
    # Use local_stub provider
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create test document
    doc = Document()
    doc.add_heading('Test Heading', level=1)
    doc.add_paragraph('This is a test paragraph with some content.')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    try:
        # Ingest with embeddings
        document = ingest_document(
            db=db_session,
            file_path=temp_file.name,
            version_tag="ingest-with-embeddings",
            source_filename="test.docx",
            create_embeddings=True
        )

        # Verify embeddings were created
        chunks = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).all()

        assert len(chunks) > 0
        for chunk in chunks:
            assert chunk.embedding is not None
            assert len(chunk.embedding) == settings.EMBEDDING_DIMENSIONS

    finally:
        os.unlink(temp_file.name)


@pytest.mark.skipif(
    not os.environ.get("OPENAI_API_KEY") or
    os.environ.get("OPENAI_API_KEY", "").startswith("sk-your"),
    reason="Requires valid OPENAI_API_KEY"
)
def test_openai_provider_integration():
    """Integration test for OpenAI provider (requires valid API key)."""
    provider = OpenAIEmbeddingProvider(
        api_key=os.environ["OPENAI_API_KEY"],
        model="text-embedding-3-small",
        dimensions=1536
    )

    embedding = provider.create_embedding("Test content for OpenAI")

    assert embedding is not None
    assert len(embedding) == 1536
    assert all(isinstance(x, float) for x in embedding)
