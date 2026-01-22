"""Browse functionality tests."""
import pytest
import os
import tempfile
from docx import Document
from sqlalchemy.orm import Session

from app.models.toolkit import ToolkitChunk
from app.services.ingestion import ingest_document
from app.services.browse import (
    browse_chunks,
    get_available_clusters,
    get_section_detail,
    search_chunks_by_text
)


def create_test_document_with_metadata() -> str:
    """Create a test DOCX file with various sections."""
    doc = Document()

    doc.add_heading('Introduction to Tools', level=1)
    doc.add_paragraph('This section introduces the key tools for productivity.')

    doc.add_heading('ChatGPT Usage', level=1)
    doc.add_paragraph('ChatGPT is a conversational AI tool that helps with various tasks.')

    doc.add_heading('Best Practices', level=1)
    doc.add_paragraph('Always validate AI outputs and use version control for your projects.')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_browse_chunks_no_filters(db_session, monkeypatch):
    """Test browsing all chunks without filters."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="browse-test-v1",
            source_filename="browse-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Browse without filters
    results = browse_chunks(db_session)

    # Should return grouped results
    assert len(results) > 0

    # Each result should have required fields
    for result in results:
        assert result.heading is not None
        assert result.excerpt is not None
        assert result.chunk_count >= 1


def test_browse_chunks_with_keyword(db_session, monkeypatch):
    """Test browsing with keyword search."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="browse-keyword-test-v1",
            source_filename="browse-keyword-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Search for "ChatGPT"
    results = browse_chunks(db_session, keyword="ChatGPT")

    # Should find results containing ChatGPT
    assert len(results) > 0

    # At least one result should mention ChatGPT
    found = False
    for result in results:
        if "ChatGPT" in result.heading or "ChatGPT" in result.excerpt:
            found = True
            break

    assert found, "Should find results mentioning ChatGPT"


def test_browse_chunks_with_cluster_filter(db_session, monkeypatch):
    """Test browsing with cluster filter."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        document = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="browse-cluster-test-v1",
            source_filename="browse-cluster-test.docx",
            create_embeddings=True
        )

        # Add metadata to chunks
        chunks = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).all()

        # Set cluster metadata on chunks
        for i, chunk in enumerate(chunks):
            chunk.metadata = {
                "cluster": "tools" if i % 2 == 0 else "practices",
                "tool_name": "TestTool"
            }
        db_session.commit()

    finally:
        os.unlink(docx_file)

    # Browse with cluster filter
    results_tools = browse_chunks(db_session, cluster="tools")
    results_practices = browse_chunks(db_session, cluster="practices")

    # Should have results for both clusters
    assert len(results_tools) > 0
    assert len(results_practices) > 0

    # Results should match the cluster filter
    for result in results_tools:
        if result.cluster:
            assert result.cluster == "tools"


def test_get_available_clusters(db_session, monkeypatch):
    """Test getting list of available clusters."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        document = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="clusters-test-v1",
            source_filename="clusters-test.docx",
            create_embeddings=True
        )

        # Add metadata with different clusters
        chunks = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).all()

        clusters_to_add = ["tools", "practices", "tools"]  # Intentional duplicate
        for i, chunk in enumerate(chunks[:3]):
            chunk.metadata = {"cluster": clusters_to_add[i]}
        db_session.commit()

    finally:
        os.unlink(docx_file)

    # Get available clusters
    clusters = get_available_clusters(db_session)

    # Should return unique, sorted clusters
    assert "tools" in clusters
    assert "practices" in clusters
    assert clusters == sorted(clusters)


def test_get_section_detail(db_session, monkeypatch):
    """Test getting full detail for a section."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="detail-test-v1",
            source_filename="detail-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Get all chunks to find a heading
    chunks = db_session.query(ToolkitChunk).all()
    assert len(chunks) > 0

    heading = chunks[0].heading

    # Get section detail
    section = get_section_detail(db_session, heading)

    assert section is not None
    assert section['heading'] == heading
    assert 'full_text' in section
    assert len(section['full_text']) > 0
    assert section['chunk_count'] > 0
    assert len(section['chunks']) == section['chunk_count']


def test_get_section_detail_nonexistent(db_session):
    """Test getting detail for non-existent section returns None."""
    section = get_section_detail(db_session, "Nonexistent Section")
    assert section is None


def test_search_chunks_by_text(db_session, monkeypatch):
    """Test text-based search in chunks."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="text-search-test-v1",
            source_filename="text-search-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Search for specific text
    results = search_chunks_by_text(db_session, "productivity")

    # Should find results
    assert len(results) > 0

    # Results should contain the search term
    found = False
    for result in results:
        if "productivity" in result['full_text'].lower():
            found = True
            break

    assert found


def test_browse_groups_by_heading(db_session, monkeypatch):
    """Test that browse groups chunks by heading."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create document with multiple chunks per heading
    doc = Document()
    doc.add_heading('Same Heading', level=1)
    doc.add_paragraph('A' * 1000)  # Force multiple chunks
    doc.add_paragraph('B' * 1000)
    doc.add_paragraph('C' * 1000)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    try:
        ingest_document(
            db=db_session,
            file_path=temp_file.name,
            version_tag="grouping-test-v1",
            source_filename="grouping-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(temp_file.name)

    # Browse results
    results = browse_chunks(db_session)

    # Should group by heading
    heading_counts = {}
    for result in results:
        heading_counts[result.heading] = result.chunk_count

    # Same heading should have multiple chunks
    assert any(count > 1 for count in heading_counts.values())


def test_browse_excerpt_truncated(db_session, monkeypatch):
    """Test that excerpts are truncated to 200 chars."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create document with long paragraph
    doc = Document()
    doc.add_heading('Long Content', level=1)
    doc.add_paragraph('X' * 500)  # Very long paragraph

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    try:
        ingest_document(
            db=db_session,
            file_path=temp_file.name,
            version_tag="excerpt-test-v1",
            source_filename="excerpt-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(temp_file.name)

    # Browse results
    results = browse_chunks(db_session)

    # Find the result with long content
    for result in results:
        if result.heading == 'Long Content':
            # Excerpt should be truncated
            assert len(result.excerpt) <= 203  # 200 + "..."
            if len(result.excerpt) > 200:
                assert result.excerpt.endswith("...")


def test_browse_no_results_empty_db(db_session):
    """Test browsing with empty database returns empty list."""
    results = browse_chunks(db_session)
    assert results == []


def test_browse_respects_is_active_flag(db_session, monkeypatch):
    """Test that browse only shows chunks from active documents."""
    from app import settings as app_settings
    from app.models.toolkit import ToolkitDocument

    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create and ingest document
    docx_file = create_test_document_with_metadata()

    try:
        document = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="active-test-v1",
            source_filename="active-test.docx",
            create_embeddings=True
        )

        # Browse should return results
        results_before = browse_chunks(db_session)
        assert len(results_before) > 0

        # Mark document as inactive
        document.is_active = False
        db_session.commit()

        # Browse should return no results
        results_after = browse_chunks(db_session)
        assert len(results_after) == 0

    finally:
        os.unlink(docx_file)
