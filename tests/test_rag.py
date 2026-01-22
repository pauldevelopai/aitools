"""RAG endpoint tests."""
import pytest
import os
import tempfile
from docx import Document
from sqlalchemy.orm import Session

from app.models.toolkit import ToolkitChunk, ChatLog
from app.services.ingestion import ingest_document
from app.services.rag import search_similar_chunks, rag_answer
from app.settings import settings


def create_test_document_with_content(content_blocks: list) -> str:
    """Create a test DOCX file with specific content."""
    doc = Document()
    for block in content_blocks:
        if block['type'] == 'heading':
            doc.add_heading(block['text'], level=1)
        else:
            doc.add_paragraph(block['text'])

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


@pytest.fixture
def ingested_document(db_session, monkeypatch):
    """Create and ingest a test document with embeddings."""
    # Use local_stub provider for testing
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    content_blocks = [
        {'type': 'heading', 'text': 'Getting Started'},
        {'type': 'paragraph', 'text': 'This toolkit helps you learn about AI tools and best practices for using them in your workflow.'},
        {'type': 'heading', 'text': 'Best Practices'},
        {'type': 'paragraph', 'text': 'Always validate AI outputs before using them in production. Use version control for all code changes.'},
        {'type': 'heading', 'text': 'Common Pitfalls'},
        {'type': 'paragraph', 'text': 'Avoid blindly trusting AI-generated code. Always review and test thoroughly.'},
    ]

    docx_file = create_test_document_with_content(content_blocks)

    try:
        # Ingest with embeddings
        document = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="test-rag-v1",
            source_filename="test-rag.docx",
            create_embeddings=True
        )

        return document

    finally:
        os.unlink(docx_file)


def test_search_returns_chunks_after_ingest(db_session, ingested_document, monkeypatch):
    """Test search endpoint returns chunks after ingestion."""
    # Use local_stub provider
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Search for content that should be in the document
    results = search_similar_chunks(
        db=db_session,
        query="What are the best practices for AI?",
        top_k=5,
        similarity_threshold=0.0
    )

    # Should return results
    assert len(results) > 0

    # Each result should have required fields
    for result in results:
        assert result.chunk_id is not None
        assert result.chunk_text is not None
        assert isinstance(result.similarity_score, float)
        assert result.similarity_score >= 0.0
        assert result.similarity_score <= 1.0


def test_search_with_similarity_threshold(db_session, ingested_document, monkeypatch):
    """Test search filters results by similarity threshold."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Search with high threshold
    results_high = search_similar_chunks(
        db=db_session,
        query="What are the best practices?",
        top_k=10,
        similarity_threshold=0.9
    )

    # Search with low threshold
    results_low = search_similar_chunks(
        db=db_session,
        query="What are the best practices?",
        top_k=10,
        similarity_threshold=0.1
    )

    # Low threshold should return more or equal results
    assert len(results_low) >= len(results_high)


def test_search_respects_top_k(db_session, ingested_document, monkeypatch):
    """Test search respects top_k parameter."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Search with top_k=2
    results = search_similar_chunks(
        db=db_session,
        query="AI tools",
        top_k=2,
        similarity_threshold=0.0
    )

    # Should return at most 2 results
    assert len(results) <= 2


def test_answer_returns_citations(db_session, ingested_document, monkeypatch):
    """Test answer endpoint returns citations with real chunk IDs."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")
    monkeypatch.setattr(settings, "RAG_SIMILARITY_THRESHOLD", 0.0)

    # Mock OpenAI call to avoid actual API call
    from app.services import rag
    original_generate = rag.generate_answer

    def mock_generate(db, query, search_results, save_to_log=True):
        # Build a mock answer using the search results
        if not search_results:
            return {
                "answer": "Not found in the toolkit",
                "citations": [],
                "similarity_scores": [],
                "refusal": True
            }

        citations = []
        for result in search_results:
            citations.append({
                "chunk_id": result.chunk_id,
                "heading": result.heading,
                "snippet": result.chunk_text[:200],
                "similarity_score": result.similarity_score,
                "document_version": result.document_version,
                "metadata": result.metadata
            })

        response = {
            "answer": "Based on the toolkit, you should validate AI outputs and use version control.",
            "citations": citations,
            "similarity_scores": [r.similarity_score for r in search_results],
            "refusal": False
        }

        if save_to_log:
            rag._save_chat_log(db, query, response)

        return response

    monkeypatch.setattr(rag, "generate_answer", mock_generate)

    # Generate answer
    result = rag_answer(
        db=db_session,
        query="What are the best practices?",
        top_k=3,
        similarity_threshold=0.0
    )

    # Should have answer
    assert result['answer'] is not None
    assert len(result['answer']) > 0

    # Should have citations
    assert len(result['citations']) > 0

    # Each citation should reference a real chunk
    for citation in result['citations']:
        assert 'chunk_id' in citation
        # Verify chunk exists in database
        chunk = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.id == citation['chunk_id']
        ).first()
        assert chunk is not None

    # Should not be a refusal
    assert result['refusal'] is False


def test_answer_saves_to_chat_log(db_session, ingested_document, monkeypatch):
    """Test answer saves Q&A to chat_logs."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Mock generate_answer to avoid OpenAI call
    from app.services import rag

    def mock_generate(db, query, search_results, save_to_log=True):
        response = {
            "answer": "Test answer",
            "citations": [{"chunk_id": "test", "heading": None, "snippet": "test",
                          "similarity_score": 0.8, "document_version": "v1", "metadata": {}}],
            "similarity_scores": [0.8],
            "refusal": False
        }
        if save_to_log:
            rag._save_chat_log(db, query, response)
        return response

    monkeypatch.setattr(rag, "generate_answer", mock_generate)

    # Count logs before
    logs_before = db_session.query(ChatLog).count()

    # Generate answer
    rag_answer(
        db=db_session,
        query="Test question",
        top_k=3
    )

    # Count logs after
    logs_after = db_session.query(ChatLog).count()

    # Should have added one log
    assert logs_after == logs_before + 1

    # Verify log content
    log = db_session.query(ChatLog).order_by(ChatLog.created_at.desc()).first()
    assert log.query == "Test question"
    assert log.answer == "Test answer"
    assert len(log.citations) > 0


def test_refusal_when_db_empty(db_session, monkeypatch):
    """Test refusal behavior when database is empty."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Database is empty (no ingested_document fixture)
    result = rag_answer(
        db=db_session,
        query="What are the best practices?",
        top_k=5
    )

    # Should refuse with standard message
    assert result['answer'] == "Not found in the toolkit"
    assert len(result['citations']) == 0
    assert result['refusal'] is True


def test_refusal_when_threshold_unmet(db_session, ingested_document, monkeypatch):
    """Test refusal behavior when similarity threshold is not met."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Search with impossible threshold
    result = rag_answer(
        db=db_session,
        query="What are the best practices?",
        top_k=5,
        similarity_threshold=1.0  # Perfect similarity required (impossible)
    )

    # Should refuse
    assert result['answer'] == "Not found in the toolkit"
    assert len(result['citations']) == 0
    assert result['refusal'] is True


def test_search_returns_ordered_by_similarity(db_session, ingested_document, monkeypatch):
    """Test search results are ordered by similarity score."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    results = search_similar_chunks(
        db=db_session,
        query="best practices for AI tools",
        top_k=5,
        similarity_threshold=0.0
    )

    # Should have results
    assert len(results) > 0

    # Should be ordered by similarity (descending)
    scores = [r.similarity_score for r in results]
    assert scores == sorted(scores, reverse=True)


def test_search_with_filters(db_session, monkeypatch):
    """Test search with metadata filters."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create document with metadata
    content_blocks = [
        {'type': 'heading', 'text': 'Tool A'},
        {'type': 'paragraph', 'text': 'Information about tool A.'},
    ]

    docx_file = create_test_document_with_content(content_blocks)

    try:
        # Ingest with embeddings
        document = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="test-filters-v1",
            source_filename="test-filters.docx",
            create_embeddings=True
        )

        # Add metadata to chunk
        chunk = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == document.id
        ).first()

        chunk.metadata = {
            "cluster": "tools",
            "tool_name": "Tool A"
        }
        db_session.commit()

        # Search with matching filter
        results_match = search_similar_chunks(
            db=db_session,
            query="tool information",
            top_k=5,
            filters={"cluster": "tools"}
        )

        # Search with non-matching filter
        results_no_match = search_similar_chunks(
            db=db_session,
            query="tool information",
            top_k=5,
            filters={"cluster": "other"}
        )

        # Matching filter should return results
        assert len(results_match) > 0

        # Non-matching filter should return fewer or no results
        assert len(results_no_match) <= len(results_match)

    finally:
        os.unlink(docx_file)


def test_answer_citation_format(db_session, ingested_document, monkeypatch):
    """Test citation format includes all required fields."""
    monkeypatch.setattr(settings, "EMBEDDING_PROVIDER", "local_stub")

    # Mock generate to avoid OpenAI call
    from app.services import rag

    def mock_generate(db, query, search_results, save_to_log=True):
        if not search_results:
            return {
                "answer": "Not found in the toolkit",
                "citations": [],
                "similarity_scores": [],
                "refusal": True
            }

        citations = []
        for result in search_results:
            citations.append({
                "chunk_id": result.chunk_id,
                "heading": result.heading,
                "snippet": result.chunk_text[:200],
                "similarity_score": result.similarity_score,
                "document_version": result.document_version,
                "metadata": result.metadata
            })

        return {
            "answer": "Test answer",
            "citations": citations,
            "similarity_scores": [r.similarity_score for r in search_results],
            "refusal": False
        }

    monkeypatch.setattr(rag, "generate_answer", mock_generate)

    result = rag_answer(
        db=db_session,
        query="What is the toolkit about?",
        top_k=3
    )

    # Should have citations
    assert len(result['citations']) > 0

    # Each citation should have required fields
    for citation in result['citations']:
        assert 'chunk_id' in citation
        assert 'snippet' in citation
        assert 'similarity_score' in citation
        # Optional fields
        assert 'heading' in citation
        assert 'document_version' in citation
        assert 'metadata' in citation
