"""Toolkit UI and feedback tests."""
import pytest
import os
import tempfile
from docx import Document
from sqlalchemy.orm import Session

from app.models.auth import User
from app.models.toolkit import ChatLog, Feedback
from app.services.auth import create_user, create_session
from app.services.ingestion import ingest_document


def create_test_user(db: Session, username: str = "testuser", email: str = "test@example.com") -> User:
    """Create a test user."""
    return create_user(db, email=email, username=username, password="password123")


def create_test_document_with_content(content_blocks: list) -> str:
    """Create a test DOCX file."""
    doc = Document()
    for block in content_blocks:
        if block['type'] == 'heading':
            doc.add_heading(block['text'], level=1)
        else:
            doc.add_paragraph(block['text'])

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_unauthenticated_redirected_to_login(client):
    """Test non-authenticated users are redirected to login."""
    response = client.get("/toolkit", follow_redirects=False)

    # Should redirect (303 or 307)
    assert response.status_code in [303, 307, 401]

    # If it's a redirect, should go to login
    if response.status_code in [303, 307]:
        assert "/login" in response.headers.get("location", "")


def test_authenticated_user_access_toolkit(client, db_session, monkeypatch):
    """Test authenticated user can access toolkit page."""
    # Use local_stub for testing
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Create session
    session = create_session(db_session, str(user.id))

    # Access toolkit with session cookie
    response = client.get(
        "/toolkit",
        cookies={"session": session.session_token}
    )

    assert response.status_code == 200
    assert b"Grounded Chat" in response.content
    assert user.username.encode() in response.content


def test_feedback_saved_and_linked_to_chat_log(db_session, monkeypatch):
    """Test feedback is saved and linked to chat log."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Create ingested document
    content_blocks = [
        {'type': 'heading', 'text': 'Test Section'},
        {'type': 'paragraph', 'text': 'This is test content for feedback testing.'},
    ]
    docx_file = create_test_document_with_content(content_blocks)

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="feedback-test-v1",
            source_filename="feedback-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Create a chat log
    from app.services.rag import rag_answer

    result = rag_answer(
        db=db_session,
        query="What is the test content?",
        user_id=str(user.id)
    )

    # Get the chat log
    chat_log = db_session.query(ChatLog).filter(
        ChatLog.user_id == user.id
    ).first()

    assert chat_log is not None
    assert chat_log.user_id == user.id

    # Create feedback
    feedback = Feedback(
        chat_log_id=chat_log.id,
        user_id=user.id,
        rating=4,
        issue_type="other",
        comment="Test comment"
    )
    db_session.add(feedback)
    db_session.commit()

    # Verify feedback was saved
    saved_feedback = db_session.query(Feedback).filter(
        Feedback.chat_log_id == chat_log.id
    ).first()

    assert saved_feedback is not None
    assert saved_feedback.chat_log_id == chat_log.id
    assert saved_feedback.user_id == user.id
    assert saved_feedback.rating == 4
    assert saved_feedback.issue_type == "other"
    assert saved_feedback.comment == "Test comment"


def test_user_sees_only_own_chat_logs(db_session, monkeypatch):
    """Test multi-user separation: users see only their own logs."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create two users
    user1 = create_test_user(db_session, username="user1", email="user1@example.com")
    user2 = create_test_user(db_session, username="user2", email="user2@example.com")

    # Create ingested document
    content_blocks = [
        {'type': 'heading', 'text': 'Test Section'},
        {'type': 'paragraph', 'text': 'This is test content for isolation testing.'},
    ]
    docx_file = create_test_document_with_content(content_blocks)

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="isolation-test-v1",
            source_filename="isolation-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Create chat logs for both users
    from app.services.rag import rag_answer

    # User 1 asks a question
    rag_answer(
        db=db_session,
        query="User 1 question",
        user_id=str(user1.id)
    )

    # User 2 asks a question
    rag_answer(
        db=db_session,
        query="User 2 question",
        user_id=str(user2.id)
    )

    # Get user 1's logs
    user1_logs = db_session.query(ChatLog).filter(
        ChatLog.user_id == user1.id
    ).all()

    # Get user 2's logs
    user2_logs = db_session.query(ChatLog).filter(
        ChatLog.user_id == user2.id
    ).all()

    # Each user should have exactly 1 log
    assert len(user1_logs) == 1
    assert len(user2_logs) == 1

    # User 1 should only see their own question
    assert user1_logs[0].query == "User 1 question"
    assert user1_logs[0].user_id == user1.id

    # User 2 should only see their own question
    assert user2_logs[0].query == "User 2 question"
    assert user2_logs[0].user_id == user2.id

    # Verify no cross-contamination
    assert user1_logs[0].id != user2_logs[0].id


def test_feedback_linked_to_correct_user(db_session, monkeypatch):
    """Test feedback is linked to the correct user's chat log."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Create ingested document
    content_blocks = [
        {'type': 'heading', 'text': 'Test Section'},
        {'type': 'paragraph', 'text': 'This is test content.'},
    ]
    docx_file = create_test_document_with_content(content_blocks)

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="feedback-link-test-v1",
            source_filename="feedback-link-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Create chat log
    from app.services.rag import rag_answer

    rag_answer(
        db=db_session,
        query="Test question",
        user_id=str(user.id)
    )

    # Get chat log
    chat_log = db_session.query(ChatLog).filter(
        ChatLog.user_id == user.id
    ).first()

    # Create feedback
    feedback = Feedback(
        chat_log_id=chat_log.id,
        user_id=user.id,
        rating=5
    )
    db_session.add(feedback)
    db_session.commit()

    # Query feedback by user
    user_feedback = db_session.query(Feedback).filter(
        Feedback.user_id == user.id
    ).all()

    assert len(user_feedback) == 1
    assert user_feedback[0].chat_log_id == chat_log.id
    assert user_feedback[0].user_id == user.id


def test_feedback_all_fields_optional_except_rating(db_session, monkeypatch):
    """Test feedback can be submitted with only rating (other fields optional)."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Create ingested document
    content_blocks = [
        {'type': 'heading', 'text': 'Test Section'},
        {'type': 'paragraph', 'text': 'This is test content.'},
    ]
    docx_file = create_test_document_with_content(content_blocks)

    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="feedback-minimal-test-v1",
            source_filename="feedback-minimal-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Create chat log
    from app.services.rag import rag_answer

    rag_answer(
        db=db_session,
        query="Test question",
        user_id=str(user.id)
    )

    # Get chat log
    chat_log = db_session.query(ChatLog).filter(
        ChatLog.user_id == user.id
    ).first()

    # Create minimal feedback (only rating)
    feedback = Feedback(
        chat_log_id=chat_log.id,
        user_id=user.id,
        rating=3
    )
    db_session.add(feedback)
    db_session.commit()

    # Verify feedback was saved
    saved_feedback = db_session.query(Feedback).filter(
        Feedback.chat_log_id == chat_log.id
    ).first()

    assert saved_feedback is not None
    assert saved_feedback.rating == 3
    assert saved_feedback.issue_type is None
    assert saved_feedback.comment is None
