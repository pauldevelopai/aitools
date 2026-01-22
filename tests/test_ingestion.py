"""Ingestion tests."""
import pytest
import os
import tempfile
from docx import Document
from sqlalchemy import text

from app.models.toolkit import ToolkitDocument, ToolkitChunk
from app.services.ingestion import parse_docx, chunk_content, ingest_document


def create_test_docx():
    """Create a temporary test DOCX file."""
    doc = Document()
    doc.add_heading('Test Heading 1', level=1)
    doc.add_paragraph('This is the first paragraph with some content.')
    doc.add_paragraph('This is the second paragraph with more content.')
    doc.add_heading('Test Heading 2', level=2)
    doc.add_paragraph('Another paragraph under the second heading.')
    doc.add_paragraph('And yet another paragraph with even more text to make it longer.')

    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_parse_docx():
    """Test DOCX parsing extracts content blocks."""
    docx_file = create_test_docx()
    try:
        content_blocks = parse_docx(docx_file)

        # Should have content blocks
        assert len(content_blocks) > 0

        # Should have headings and paragraphs
        types = [block['type'] for block in content_blocks]
        assert 'heading' in types
        assert 'paragraph' in types

        # Should have text content
        assert all('text' in block for block in content_blocks)
        assert all(block['text'] for block in content_blocks)

    finally:
        os.unlink(docx_file)


def test_chunk_content():
    """Test content chunking with overlap."""
    content_blocks = [
        {'type': 'heading', 'text': 'Chapter 1', 'heading': 'Chapter 1'},
        {'type': 'paragraph', 'text': 'A' * 500, 'heading': 'Chapter 1'},
        {'type': 'paragraph', 'text': 'B' * 600, 'heading': 'Chapter 1'},
        {'type': 'heading', 'text': 'Chapter 2', 'heading': 'Chapter 2'},
        {'type': 'paragraph', 'text': 'C' * 700, 'heading': 'Chapter 2'},
    ]

    chunks = chunk_content(content_blocks, target_size=1000, overlap=150)

    # Should create chunks
    assert len(chunks) > 0

    # Each chunk should have required fields
    for chunk in chunks:
        assert 'chunk_text' in chunk
        assert 'chunk_index' in chunk
        assert 'heading' in chunk
        assert 'metadata' in chunk

        # Chunks should not be empty
        assert len(chunk['chunk_text']) > 0

        # Chunks should be roughly target size
        assert len(chunk['chunk_text']) <= 2000  # Allow some flexibility


def test_ingest_document_creates_chunks(db_session):
    """Test document ingestion creates non-empty chunks (MUST PASS ALWAYS)."""
    docx_file = create_test_docx()
    try:
        # Ingest without embeddings (so test doesn't require API key)
        doc = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="test-v1",
            source_filename="test.docx",
            create_embeddings=False
        )

        # Document should be created
        assert doc is not None
        assert doc.version_tag == "test-v1"
        assert doc.source_filename == "test.docx"

        # Chunks should be created (CRITICAL TEST)
        assert doc.chunk_count > 0, "No chunks created! Ingestion must create chunks."

        # Verify chunks exist in database
        chunks = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == doc.id
        ).all()

        assert len(chunks) == doc.chunk_count
        assert len(chunks) > 0, "toolkit_chunks table is empty after ingestion!"

        # All chunks must have non-empty text
        for chunk in chunks:
            assert chunk.chunk_text is not None
            assert len(chunk.chunk_text) > 0, f"Chunk {chunk.id} has empty chunk_text!"
            assert chunk.chunk_index >= 0

    finally:
        os.unlink(docx_file)


def test_ingest_duplicate_version_fails(db_session):
    """Test ingesting duplicate version tag fails."""
    docx_file = create_test_docx()
    try:
        # First ingest should succeed
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="duplicate-test",
            source_filename="test.docx",
            create_embeddings=False
        )

        # Second ingest with same version should fail
        with pytest.raises(ValueError, match="already exists"):
            ingest_document(
                db=db_session,
                file_path=docx_file,
                version_tag="duplicate-test",
                source_filename="test2.docx",
                create_embeddings=False
            )

    finally:
        os.unlink(docx_file)


def test_chunks_have_proper_structure(db_session):
    """Test chunks have all required fields."""
    docx_file = create_test_docx()
    try:
        doc = ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="structure-test",
            source_filename="test.docx",
            create_embeddings=False
        )

        chunks = db_session.query(ToolkitChunk).filter(
            ToolkitChunk.document_id == doc.id
        ).all()

        for chunk in chunks:
            # Required fields
            assert chunk.id is not None
            assert chunk.document_id == doc.id
            assert chunk.chunk_text is not None
            assert isinstance(chunk.chunk_index, int)

            # Optional fields should be present but can be None
            assert hasattr(chunk, 'heading')
            assert hasattr(chunk, 'metadata')
            assert hasattr(chunk, 'embedding')

    finally:
        os.unlink(docx_file)
